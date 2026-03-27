# 文档读取模块
# 用于读取和分析项目文档内容

import re
import os
from pathlib import Path
from typing import Dict, List, Any, Optional

class DocumentReader:
    """文档读取器，支持多种格式"""

    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.md', '.txt', '.doc']

    def read_document(self, file_path: str) -> Dict[str, Any]:
        """
        读取文档并返回结构化数据

        Args:
            file_path: 文档路径

        Returns:
            包含文档结构和内容的字典
        """
        file_path = Path(file_path)

        # 检查文件是否存在
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        # 检查文件格式
        file_extension = file_path.suffix.lower()
        if file_extension not in self.supported_formats:
            raise ValueError(f"不支持的文件格式: {file_extension}")

        # 根据文件类型读取
        if file_extension == '.pdf':
            content = self._read_pdf(file_path)
        elif file_extension in ['.doc', '.docx']:
            content = self._read_word(file_path)
        elif file_extension == '.md':
            content = self._read_markdown(file_path)
        else:
            content = self._read_text(file_path)

        # 解析文档结构
        structured_data = self._parse_document_structure(content)

        return {
            'file_path': str(file_path),
            'file_type': file_extension,
            'content': content,
            'structured_data': structured_data,
            'basic_info': self._extract_basic_info(structured_data)
        }

    def _read_pdf(self, file_path: Path) -> str:
        """读取PDF文件"""
        try:
            import PyPDF2
            content = ""
            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    content += page.extract_text() + "\n"
            return content
        except ImportError:
            raise ImportError("需要安装PyPDF2库来读取PDF文件: pip install PyPDF2")

    def _read_word(self, file_path: Path) -> str:
        """读取Word文档"""
        try:
            import docx
            doc = docx.Document(file_path)
            content = ""

            # 读取段落
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"

            # 读取表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += f"[表格] {cell.text} | "
                    content += "\n"

            return content.strip()
        except ImportError:
            raise ImportError("需要安装python-docx库来读取Word文件: pip install python-docx")
        except Exception as e:
            raise Exception(f"读取Word文件失败: {str(e)}")

    def _read_markdown(self, file_path: Path) -> str:
        """读取Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"读取Markdown文件失败: {str(e)}")

    def _read_text(self, file_path: Path) -> str:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            raise Exception(f"读取文本文件失败: {str(e)}")

    def _parse_document_structure(self, content: str) -> Dict[str, Any]:
        """解析文档结构"""
        # 按段落分割内容
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]

        # 检测可能的标题
        headers = []
        body_content = []

        for para in paragraphs:
            # 简单的标题检测
            if (len(para) < 100 and
                any(keyword in para for keyword in ['项目', '职责', '任务', '成果', '目标', '背景', '技术', '开发'])):
                headers.append(para)
            else:
                body_content.append(para)

        # 提取关键信息
        project_info = self._extract_project_info(content)
        responsibilities = self._extract_responsibilities(content)
        achievements = self._extract_achievements(content)
        technologies = self._extract_technologies(content)

        return {
            'headers': headers,
            'paragraphs': paragraphs,
            'body_content': body_content,
            'project_info': project_info,
            'responsibilities': responsibilities,
            'achievements': achievements,
            'technologies': technologies
        }

    def _extract_basic_info(self, structured_data: Dict[str, Any]) -> Dict[str, Any]:
        """提取基本信息"""
        content = ' '.join(structured_data['paragraphs'])

        # 提取项目名称
        project_patterns = [
            r'项目[：:](.+)',
            r'(.+?)项目',
            r'(.+?)系统',
            r'(.+?)平台',
            r'(.+?)应用'
        ]

        project_name = None
        for pattern in project_patterns:
            match = re.search(pattern, content)
            if match:
                project_name = match.group(1)
                break

        # 提取时间范围
        time_patterns = [
            r'(\d{4})\.(\d{1,2})-(\d{4})\.(\d{1,2})',
            r'(\d{4})年(\d{1,2})月-(\d{4})年(\d{1,2})月',
            r'(\d{4})\.(\d{1,2})至今',
            r'(\d{4})年(\d{1,2})月至今'
        ]

        time_range = None
        for pattern in time_patterns:
            match = re.search(pattern, content)
            if match:
                time_range = match.group(0)
                break

        # 提取角色/职位
        role_patterns = [
            r'(.+?)负责人',
            r'(.+?)开发者',
            r'(.+?)工程师',
            r'(.+?)分析师',
            r'(.+?)经理'
        ]

        role = None
        for pattern in role_patterns:
            match = re.search(pattern, content)
            if match:
                role = match.group(1)
                break

        return {
            'project_name': project_name,
            'time_range': time_range,
            'role': role,
            'content_length': len(content)
        }

    def _extract_project_info(self, content: str) -> List[str]:
        """提取项目信息"""
        project_info = []

        # 查找项目描述
        sentences = re.split(r'[。！？\n]', content)
        for sentence in sentences:
            sentence = sentence.strip()
            if (len(sentence) > 10 and
                any(keyword in sentence for keyword in ['项目', '系统', '平台', '应用', '目标'])):
                project_info.append(sentence)

        return project_info[:3]

    def _extract_responsibilities(self, content: str) -> List[str]:
        """提取职责描述"""
        responsibilities = []

        # 查找职责相关内容
        responsibility_keywords = ['负责', '参与', '协助', '完成', '开发', '设计', '分析', '测试']

        sentences = re.split(r'[。！？\n]', content)
        for sentence in sentences:
            sentence = sentence.strip()
            if (len(sentence) > 10 and
                any(keyword in sentence for keyword in responsibility_keywords)):
                responsibilities.append(sentence)

        return responsibilities[:8]

    def _extract_achievements(self, content: str) -> List[str]:
        """提取成果"""
        achievements = []

        # 查找包含数字的成果
        sentences = re.split(r'[。！？\n]', content)
        for sentence in sentences:
            sentence = sentence.strip()
            if (len(sentence) > 10 and
                (re.search(r'\d+%', sentence) or  # 包含百分比
                 re.search(r'\d+万', sentence) or  # 包含万
                 re.search(r'\d+亿', sentence) or  # 包含亿
                 re.search(r'\d+倍', sentence) or  # 包含倍
                 (re.search(r'\d+', sentence) and
                  any(word in sentence for word in ['提升', '降低', '增加', '节省'])))):
                achievements.append(sentence)

        return achievements[:6]

    def _extract_technologies(self, content: str) -> List[str]:
        """提取使用的技术"""
        tech_keywords = [
            'Java', 'Python', 'JavaScript', 'React', 'Vue', 'Angular',
            'Spring', 'Django', 'Flask', 'Node.js', 'MySQL', 'Redis',
            'MongoDB', 'Docker', 'Kubernetes', 'Git', 'Jenkins'
        ]

        found_tech = []
        for tech in tech_keywords:
            if tech in content:
                found_tech.append(tech)

        return found_tech

def analyze_project_document(file_path: str) -> Dict[str, Any]:
    """
    分析项目文档的主函数

    Args:
        file_path: 文档路径

    Returns:
        包含分析结果的结构化数据
    """
    reader = DocumentReader()
    document_data = reader.read_document(file_path)

    # 进一步分析和提炼关键信息
    content = document_data['content']

    # 提取STAR法则要素
    star_elements = extract_star_elements(content)

    # 生成四字词语
    four_char_words = generate_four_char_words(star_elements)

    # 量化成果
    quantified_results = quantify_achievements(document_data['structured_data']['achievements'])

    # 更新分析结果
    document_data['analysis_result'] = {
        'star_elements': star_elements,
        'four_char_words': four_char_words,
        'quantified_results': quantified_results
    }

    return document_data

def extract_star_elements(content: str) -> Dict[str, str]:
    """提取STAR法则要素"""
    return {
        'situation': extract_situation(content),
        'task': extract_task(content),
        'action': extract_action(content),
        'result': extract_result(content)
    }

def extract_situation(content: str) -> str:
    """提取背景情况"""
    # 查找背景描述
    background_keywords = ['背景', '目标', '需求', '问题', '挑战']

    sentences = re.split(r'[。！？\n]', content)
    for sentence in sentences:
        sentence = sentence.strip()
        if (len(sentence) > 20 and
            any(keyword in sentence for keyword in background_keywords)):
            return sentence

    return "负责项目的设计、开发和优化工作"

def extract_task(content: str) -> str:
    """提取任务"""
    # 查找任务描述
    task_keywords = ['任务', '职责', '目标', '需求']

    sentences = re.split(r'[。！？\n]', content)
    for sentence in sentences:
        sentence = sentence.strip()
        if (len(sentence) > 15 and
            any(keyword in sentence for keyword in task_keywords)):
            return sentence

    return "完成核心功能开发，优化系统性能"

def extract_action(content: str) -> str:
    """提取行动"""
    # 查找行动描述
    action_verbs = ['设计', '开发', '优化', '实现', '测试', '部署', '维护']

    sentences = re.split(r'[。！？\n]', content)
    for sentence in sentences:
        sentence = sentence.strip()
        if (len(sentence) > 15 and
            any(verb in sentence for verb in action_verbs)):
            return sentence

    return "使用相关技术栈进行开发，解决关键技术问题"

def extract_result(content: str) -> str:
    """提取成果"""
    # 查找成果描述
    achievements = [a for a in content.split('。') if len(a.strip()) > 10 and
                   any(word in a for word in ['提升', '降低', '增加', '节省', '优化', '完成'])]

    if achievements:
        return achievements[0].strip()

    return "提高了系统性能，改善了用户体验"

def generate_four_char_words(star_elements: Dict[str, str]) -> List[str]:
    """生成四字词语"""
    # 基于STAR要素生成四字词语
    word_map = {
        'situation': ['需求分析', '项目规划', '背景调研'],
        'task': ['任务开发', '功能实现', '目标达成'],
        'action': ['核心开发', '技术攻关', '系统优化'],
        'result': ['效率提升', '质量保证', '价值创造']
    }

    result = []
    for element, words in word_map.items():
        result.append(words[0])  # 取第一个词作为示例

    return result

def quantify_achievements(achievements: List[str]) -> List[str]:
    """量化成果"""
    quantified = []
    for achievement in achievements:
        if '提升' in achievement or '降低' in achievement:
            quantified.append(achievement)
        elif re.search(r'\d+', achievement):
            quantified.append(achievement)

    return quantified[:3]

# 使用示例
if __name__ == "__main__":
    # 测试文档读取
    test_file = "test_project.docx"
    if os.path.exists(test_file):
        try:
            data = analyze_project_document(test_file)
            print("基本信息:")
            print(f"项目名称: {data['basic_info']['project_name']}")
            print(f"角色: {data['basic_info']['role']}")
            print(f"时间: {data['basic_info']['time_range']}")
            print("\n分析结果:")
            print(f"四字词语: {data['analysis_result']['four_char_words']}")
        except Exception as e:
            print(f"读取文档失败: {str(e)}")
    else:
        print("测试文件不存在")